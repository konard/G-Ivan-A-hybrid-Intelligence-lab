#!/usr/bin/env python3
"""Extract text from docx and xlsx files into per-file .txt artifacts."""
import os
import sys
import zipfile
from pathlib import Path

import docx
import openpyxl

SRC = Path("files")
OUT = Path("text")
OUT.mkdir(exist_ok=True)


def extract_docx(path: Path) -> str:
    try:
        d = docx.Document(str(path))
    except Exception as exc:
        return f"[ERROR python-docx: {exc}]"
    chunks = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            chunks.append(t)
    for ti, tbl in enumerate(d.tables):
        chunks.append(f"\n[TABLE {ti+1}]")
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def extract_xlsx(path: Path) -> str:
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    out = []
    for sn in wb.sheetnames:
        ws = wb[sn]
        out.append(f"\n[SHEET: {sn}]")
        for row in ws.iter_rows(values_only=True):
            line = " | ".join("" if v is None else str(v).strip() for v in row)
            if line.replace("|", "").strip():
                out.append(line)
    return "\n".join(out)


def main() -> None:
    files = sorted(SRC.iterdir(), key=lambda p: int(p.stem))
    for f in files:
        idx = f.stem
        ext = f.suffix.lower()
        target = OUT / f"{idx}.txt"
        if target.exists():
            continue
        if ext == ".docx":
            content = extract_docx(f)
        elif ext == ".xlsx":
            content = extract_xlsx(f)
        else:
            # .doc handled separately via libreoffice
            continue
        target.write_text(content, encoding="utf-8")
        print(f"wrote {target} ({len(content)} chars)")


if __name__ == "__main__":
    main()
